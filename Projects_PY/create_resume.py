from docx import Document

# Create a new Document
doc = Document()

# Add the resume content
doc.add_heading('Omar Mohamed', 0)
doc.add_heading('Software Developer', level=1)

# Personal Information
doc.add_heading('Personal Information', level=2)
doc.add_paragraph('Address:\nCairo, Egypt')
doc.add_paragraph('Phone:\n+201003765335')
doc.add_paragraph('Email:\nomarmosharawi@gmail.com')
doc.add_paragraph('GitHub:\nhttps://github.com/omarmosharawi')

# Languages and Technologies
doc.add_heading('Languages and Technologies', level=2)
doc.add_paragraph('HTML, CSS\nJavaScript\nJava, Python\nC++, C#\nPHP, R\nRuby, SQL')

# Applications and Tools
doc.add_heading('Applications and Tools', level=2)
doc.add_paragraph('GitHub, Visual Studio\nPyCharm, CodeBlocks\nNetBeans, Oracle\nPostgreSQL, MongoDB\nphpMyAdmin, MySQL')

# Hobbies
doc.add_heading('Hobbies', level=2)
doc.add_paragraph('Football\nSelf-development\nTraveling')

# Languages
doc.add_heading('Languages', level=2)
doc.add_paragraph('Arabic\nEnglish')

# Profile
doc.add_heading('Profile', level=2)
doc.add_paragraph("A diligent university student passionate about technology, proficient in application development using Python. "
                  "I have completed several training projects and aspire to work as a software developer in a motivating environment within "
                  "a professional team to prove myself and gain more experience in the field.")

# Skills
doc.add_heading('Skills', level=2)
doc.add_paragraph("Front-end Development:\nProficient in creating interactive and user-friendly web pages using HTML, CSS, JavaScript, and "
                  "the Bootstrap framework for responsive designs compatible with all screens.")
doc.add_paragraph("Python:\nPractical experience in developing applications using Python, including a project to create an application that allows "
                  "users to perform basic operations on a school database system. The application uses SQLite3 as the database engine and provides a "
                  "command-line interface for user interaction.")
doc.add_paragraph("Troubleshooting:\nAble to analyze errors, identify their causes, fix them, and provide effective solutions. I rely on a methodical "
                  "approach in analysis and seek information from online sources and trusted technical communities such as Hsoub Encyclopedia, Stack Overflow, and W3Schools.")
doc.add_paragraph("Professional Development:\nKeeping up with technological advancements and the latest market trends. I completed a computer science course "
                  "from Hsoub Academy and started studying a Python application development course and an artificial intelligence course.")

# Education
doc.add_heading('Education', level=2)
doc.add_paragraph("Canadian International College (CIC), Egypt\nBachelor's Degree, BTech\nSeptember 2022 - September 2026")
doc.add_paragraph("Ibn Khaldun Military Secondary School, Egypt\nGeneral Secondary Education\nJuly 2019 - July 2022")

# Certifications
doc.add_heading('Certifications', level=2)
doc.add_paragraph("Computer Science Course\nHsoub Academy\nSeptember 2023")

# Save the document
doc.save("Omar_Mohamed_Resume_English.docx")
